import os
import uuid
import re
import random
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image
from sqlalchemy.orm import Session
from ..models import Upload, Job, Screenshot
from ..config import settings


class ComposerService:
    """Service for composing final DOCX reports with embedded screenshots"""
    
    def __init__(self):
        pass
    
    async def compose_report(
        self, 
        upload_id: int, 
        screenshot_order: List[int], 
        db: Session
    ) -> Dict[str, Any]:
        """
        Update the original document with embedded screenshots
        
        Args:
            upload_id: ID of the uploaded document
            screenshot_order: Ordered list of task IDs for screenshots
            db: Database session
            
        Returns:
            Dict with report_path, filename, and download_url
        """
        
        # Get the original upload
        upload = db.query(Upload).filter(Upload.id == upload_id).first()
        if not upload:
            raise ValueError("Upload not found")
        
        # Get jobs with screenshots in the specified order
        # screenshot_order can contain either job IDs or task IDs
        # We'll try job IDs first, then fall back to task IDs
        jobs_with_screenshots = []
        
        processed_jobs = set()

        def append_job_screens(job: Job):
            screenshots = db.query(Screenshot).filter(Screenshot.job_id == job.id).order_by(Screenshot.created_at).all()
            has_screens = False
            for screenshot in screenshots:
                if screenshot and os.path.exists(screenshot.file_path):
                    jobs_with_screenshots.append({
                        'job': job,
                        'screenshot': screenshot,
                        'task_id': job.task_id
                    })
                    has_screens = True
            if has_screens:
                processed_jobs.add(job.id)

        if not screenshot_order:
            all_jobs = db.query(Job).filter(Job.upload_id == upload_id, Job.status == 'completed').all()
            for job in all_jobs:
                append_job_screens(job)
        else:
            for order_id in screenshot_order:
                job = db.query(Job).filter(Job.id == order_id, Job.upload_id == upload_id).first()
                if not job:
                    job = db.query(Job).filter(Job.task_id == order_id, Job.upload_id == upload_id).first()
                if job:
                    append_job_screens(job)
            remaining_jobs = db.query(Job).filter(
                Job.upload_id == upload_id,
                Job.status == 'completed'
            ).all()
            for job in remaining_jobs:
                if job.id not in processed_jobs:
                    append_job_screens(job)
        
        # Load the original document
        if not os.path.exists(upload.file_path):
            raise ValueError("Original document not found")
        
        # Create a copy of the original document
        doc = Document(upload.file_path)
        
        # Set default font to Times New Roman for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        
        # Insert screenshots at the end of the document
        await self._add_screenshots_to_document(doc, jobs_with_screenshots)
        
        # Save the updated document
        report_filename = f"{upload.filename.replace('.docx', '')}_with_screenshots_{uuid.uuid4().hex[:8]}.docx"
        report_path = os.path.join(settings.REPORT_DIR, report_filename)
        
        doc.save(report_path)
        
        return {
            "report_path": report_path,
            "filename": report_filename,
            "download_url": f"/reports/{report_filename}"
        }
    
    async def _add_screenshots_to_document(self, doc: Document, jobs_with_screenshots: List[Dict]):
        """Add screenshots to the original document with clean formatting"""
        
        # Add header with varied text to prevent AI detection
        header_text = self._get_varied_header()
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(header_text)
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()  # Add spacing after header
        
        # Add all screenshots in clean format at the end
        await self._add_formatted_screenshots(doc, jobs_with_screenshots)
    
    def _get_varied_header(self) -> str:
        """Generate a varied header to prevent AI detection"""
        headers = [
            "Lab programs with output",
            "Program outputs and execution results",
            "Laboratory exercises with results",
            "Code execution outputs",
            "Program implementations with screenshots",
            "Lab assignment outputs",
            "Execution results for programs",
            "Program outputs and screenshots"
        ]
        return random.choice(headers)
    
    async def _add_formatted_screenshots(self, doc: Document, jobs_with_screenshots: List[Dict]):
        """Add screenshots in clean formatted way: number, question, image, description"""
        
        if not jobs_with_screenshots:
            return
        
        grouped_map: Dict[int, Dict[str, Any]] = {}
        grouped_list: List[Dict[str, Any]] = []

        for item in jobs_with_screenshots:
            job = item['job']
            job_id = job.id
            if job_id not in grouped_map:
                grouped_map[job_id] = {"job": job, "screens": []}
                grouped_list.append(grouped_map[job_id])
            grouped_map[job_id]["screens"].append(item["screenshot"])

        grouped_list.sort(key=lambda entry: entry["job"].task_id)

        for index, group in enumerate(grouped_list, start=1):
            job = group["job"]
            question_text = job.question_text or f"Task {job.task_id}"

            number_para = doc.add_paragraph()
            number_run = number_para.add_run(f"{index}) ")
            number_run.font.name = 'Times New Roman'
            number_run.font.size = Pt(12)
            number_run.font.bold = True
            question_run = number_para.add_run(question_text.strip())
            question_run.font.name = 'Times New Roman'
            question_run.font.size = Pt(12)
            question_run.font.bold = False
            doc.add_paragraph()

            ordered_screens = self._order_screens_for_job(group["screens"])
            step = 1

            for entry in ordered_screens:
                screen = entry["screen"]
                label = entry["label"]
                category = entry["category"]

                self._add_step_heading(doc, f"{step}) {label}")
                if screen and os.path.exists(screen.file_path):
                    await self._add_screenshot_image(doc, screen.file_path)
                step += 1

                if category == "code":
                    description = self._generate_image_description(question_text, job, index)
                    self._add_description_line(doc, step, description)
                    step += 1

            doc.add_paragraph()
    
    def _generate_image_description(self, question_text: str, job, index: int) -> str:
        """Generate a precise two-line description for the image based on question and execution output"""
        # Extract key words from question for context
        question_lower = question_text.lower()
        
        # Get output text if available for more precise description
        output_text = getattr(job, 'output_text', '') or ''
        output_preview = output_text[:100].strip() if output_text else ''
        
        # Determine program type from question
        if any(word in question_lower for word in ['factorial', 'fibonacci', 'recursion']):
            program_type = "recursive algorithm"
        elif any(word in question_lower for word in ['sort', 'array', 'list']):
            program_type = "array operation"
        elif any(word in question_lower for word in ['pattern', 'star', 'triangle']):
            program_type = "pattern printing"
        elif any(word in question_lower for word in ['function', 'def']):
            program_type = "function implementation"
        elif any(word in question_lower for word in ['string', 'palindrome', 'reverse']):
            program_type = "string manipulation"
        elif any(word in question_lower for word in ['class', 'object', 'inheritance', 'oop']):
            program_type = "object-oriented program"
        elif any(word in question_lower for word in ['loop', 'while', 'for']):
            program_type = "loop-based program"
        else:
            program_type = "program"
        
        # Generate precise descriptions based on output
        if output_preview:
            # Use output preview to make description more specific
            if any(word in output_preview.lower() for word in ['error', 'exception', 'failed']):
                status_desc = "execution with error output"
            elif any(word in output_preview.lower() for word in ['success', 'completed', 'done']):
                status_desc = "successful execution"
            else:
                status_desc = "program execution"
            
            # Create precise description
            descriptions = [
                f"Screenshot showing {program_type} {status_desc}.\nTerminal output displays: {output_preview[:80]}...",
                f"Execution result of {program_type} implementation.\nOutput captured: {output_preview[:80]}...",
                f"Program screenshot for {program_type} demonstrating {status_desc}.\nCode output: {output_preview[:80]}..."
            ]
        else:
            # Fallback to generic but still precise descriptions
            descriptions = [
                f"Screenshot showing {program_type} execution result.\nTerminal output displaying program output and execution status.",
                f"Execution output of {program_type} implementation.\nScreenshot captures code output and terminal response.",
                f"Program result screenshot for {program_type}.\nOutput demonstrates successful program execution."
            ]
        
        return random.choice(descriptions)

    def _order_screens_for_job(self, screens: List[Any]) -> List[Dict[str, Any]]:
        buckets: Dict[str, List[Dict[str, Any]]] = {
            "txt": [],
            "pkl": [],
            "other_file": [],
            "code": []
        }
        for screen in screens:
            category, label = self._classify_screen_for_doc(screen)
            buckets[category].append({
                "screen": screen,
                "label": label,
                "category": category
            })
        ordered: List[Dict[str, Any]] = []
        for key in ["txt", "pkl", "other_file", "code"]:
            ordered.extend(buckets[key])
        return ordered

    def _classify_screen_for_doc(self, screen: Any) -> Tuple[str, str]:
        file_path = getattr(screen, "file_path", "") if screen else ""
        if self._is_file_preview_path(file_path):
            label = self._extract_file_label(file_path)
            ext = self._infer_preview_extension(file_path)
            if ext == ".pkl":
                return "pkl", f"{label or 'File'} (.pkl file)"
            if ext in {".txt", ".csv", ".log", ".dat"}:
                return "txt", f"{label or 'File'} ({ext} file)"
            return "other_file", f"{label or 'File'} ({ext or 'file'})"
        return "code", "Python code"

    def _is_file_preview_path(self, file_path: str) -> bool:
        if not file_path:
            return False
        return os.path.basename(file_path).startswith("file_")

    def _infer_preview_extension(self, file_path: str) -> str:
        if not file_path or not self._is_file_preview_path(file_path):
            return ""
        base = os.path.basename(file_path)
        core = os.path.splitext(base)[0]
        if core.startswith("file_"):
            core = core[5:]
        if "_" in core:
            core = core.rsplit("_", 1)[0]
        _, ext = os.path.splitext(core)
        return ext.lower()

    def _extract_file_label(self, file_path: str) -> Optional[str]:
        if not file_path:
            return None
        base = os.path.basename(file_path)
        if not base.startswith("file_"):
            return None
        core = os.path.splitext(base)[0]
        if core.startswith("file_"):
            core = core[5:]
        if "_" in core:
            core = core.rsplit("_", 1)[0]
        return core.replace("_", " ").strip() or None

    def _add_step_heading(self, doc: Document, text: str):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True

    def _add_description_line(self, doc: Document, step: int, description: str):
        para = doc.add_paragraph()
        run = para.add_run(f"{step}) Description: {description}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = False
    
    async def _add_screenshot_image(self, doc: Document, image_path: str):
        """Add screenshot image centered in the document"""
        try:
            # Get image dimensions
            with Image.open(image_path) as img:
                width, height = img.size
            
            # Calculate size to fit within document (max 5.5 inches width)
            max_width = Inches(5.5)
            max_height = Inches(4)
            
            # Calculate aspect ratio
            aspect_ratio = width / height
            
            if width > max_width:
                new_width = max_width
                new_height = new_width / aspect_ratio
            else:
                new_width = Inches(width / 100)  # Convert pixels to inches (approximate)
                new_height = Inches(height / 100)
            
            if new_height > max_height:
                new_height = max_height
                new_width = new_height * aspect_ratio
            
            # Add image paragraph
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image to paragraph
            run = img_para.runs[0] if img_para.runs else img_para.add_run()
            doc.add_picture(image_path, width=new_width, height=new_height)
            
            # Add spacing after image
            doc.add_paragraph()
            
        except Exception as e:
            # If image fails to load, skip it
            pass
    
    async def _add_image_description(self, doc: Document, description: str):
        """Add image description with 'Description:' in bold, followed by the description text"""
        desc_para = doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add "Description:" in bold
        desc_label_run = desc_para.add_run("Description: ")
        desc_label_run.font.name = 'Times New Roman'
        desc_label_run.font.size = Pt(11)
        desc_label_run.font.bold = True
        desc_label_run.font.italic = False
        
        # Split description into lines if it contains \n
        lines = description.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                desc_para.add_run('\n')
            desc_run = desc_para.add_run(line.strip())
            desc_run.font.name = 'Times New Roman'
            desc_run.font.size = Pt(11)
            desc_run.font.bold = False
            desc_run.font.italic = False
        
        # Add spacing after description
        doc.add_paragraph()
    
    def _find_question_end_index(self, paragraphs, start_index: int) -> int:
        """Find where a question ends (next numbered item or section)"""
        for i in range(start_index + 1, len(paragraphs)):
            text = paragraphs[i].text.strip()
            if not text:
                continue
            
            # Check if this is a new numbered question
            if self._find_question_pattern(text):
                return i
            
            # Check if this is a new section (like "C. Questions:", "A. Theory:", etc.)
            if re.match(r'^[A-Z]\.\s+[A-Z]', text):
                return i
        
        # If no end found, return the next paragraph index
        return start_index + 1
    
    def _extract_task_number(self, description: str) -> Optional[int]:
        """Extract task number from description"""
        if not description:
            return None
        
        # Look for patterns like "Task 1", "Question 1", "1.", etc.
        patterns = [
            r'Task\s+(\d+)',
            r'Question\s+(\d+)',
            r'Problem\s+(\d+)',
            r'Exercise\s+(\d+)',
            r'^(\d+)\.',
            r'^(\d+)\)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, description, re.IGNORECASE)
            if match:
                return int(match.group(1))
        
        return None
    
    def _find_question_pattern(self, text: str) -> Optional[int]:
        """Find question pattern in text and return task number"""
        patterns = [
            r'Question\s+(\d+)',
            r'Task\s+(\d+)',
            r'Problem\s+(\d+)',
            r'Exercise\s+(\d+)',
            r'^(\d+)\.',                    # "1.Write a Python program..."
            r'^(\d+)\)',                    # "1) Write a Python program..."
            r'Q(\d+)',                      # "Q1", "Q2"
            r'T(\d+)',                      # "T1", "T2"
            r'^\s*(\d+)\.\s*Write',         # "1. Write a Python program..."
            r'^\s*(\d+)\)\s*Write',         # "1) Write a Python program..."
            r'^\s*(\d+)\.\s*[A-Z]',         # "1. Demonstrate", "2. Calculate"
            r'^\s*(\d+)\)\s*[A-Z]'          # "1) Demonstrate", "2) Calculate"
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                return int(match.group(1))
        
        return None
    
    async def _insert_screenshot_after_paragraph(self, doc: Document, para_index: int, item: Dict):
        """Insert screenshot and image description below the question (no answer section)"""
        
        job = item['job']
        screenshot = item['screenshot']
        
        # Add the screenshot if available
        if screenshot and os.path.exists(screenshot.file_path):
            # Use job question_text as caption
            caption = f"Output for Task {job.task_id}: {job.question_text[:50]}..." if len(job.question_text) > 50 else f"Output for Task {job.task_id}"
            await self._add_screenshot_with_description(doc, screenshot.file_path, caption)
    
    async def _add_screenshot_with_description(self, doc: Document, image_path: str, image_description: Optional[str]):
        """Add screenshot with 1-2 line description below it (legacy method, kept for compatibility)"""
        await self._add_screenshot_image(doc, image_path)
        if image_description:
            await self._add_image_description(doc, image_description)
    
    async def _add_screenshot_only_clean(self, doc: Document, image_path: str, caption: str):
        """Add only the screenshot with minimal caption - no extra text or formatting"""
        
        try:
            # Get image dimensions
            with Image.open(image_path) as img:
                width, height = img.size
            
            # Calculate size to fit within document (max 6 inches width)
            max_width = Inches(6)
            max_height = Inches(4)
            
            # Calculate aspect ratio
            aspect_ratio = width / height
            
            if width > max_width:
                new_width = max_width
                new_height = new_width / aspect_ratio
            else:
                new_width = Inches(width / 100)  # Convert pixels to inches (approximate)
                new_height = Inches(height / 100)
            
            if new_height > max_height:
                new_height = max_height
                new_width = new_height * aspect_ratio
            
            # Add image
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image to paragraph
            doc.add_picture(image_path, width=new_width, height=new_height)
            
            # Add minimal caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            
        except Exception as e:
            # If image fails to load, add minimal error message
            error_para = doc.add_paragraph()
            error_para.add_run(f"Screenshot unavailable")
            error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    async def _add_screenshot(self, doc: Document, image_path: str, caption: str):
        """Add a screenshot to the document with proper sizing and caption"""
        
        try:
            # Get image dimensions
            with Image.open(image_path) as img:
                width, height = img.size
            
            # Calculate size to fit within document (max 6 inches width)
            max_width = Inches(6)
            max_height = Inches(4)
            
            # Calculate aspect ratio
            aspect_ratio = width / height
            
            if width > max_width:
                new_width = max_width
                new_height = new_width / aspect_ratio
            else:
                new_width = Inches(width / 100)  # Convert pixels to inches (approximate)
                new_height = Inches(height / 100)
            
            if new_height > max_height:
                new_height = max_height
                new_width = new_height * aspect_ratio
            
            # Add image
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            # Add image to paragraph
            doc.add_picture(image_path, width=new_width, height=new_height)
            
            # Add caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            
        except Exception as e:
            # If image fails to load, add error message
            error_para = doc.add_paragraph()
            error_para.add_run(f"Error loading screenshot: {str(e)}")
            error_para.runs[0].font.color.rgb = self._get_status_color('failed')
    
    def _get_status_color(self, status: str):
        """Get RGB color for status"""
        from docx.shared import RGBColor
        
        if status == 'completed':
            return RGBColor(0, 128, 0)  # Green
        elif status == 'failed':
            return RGBColor(255, 0, 0)  # Red
        else:
            return RGBColor(128, 128, 128)  # Gray


# Create singleton instance
composer_service = ComposerService()
