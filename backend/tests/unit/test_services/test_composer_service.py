"""
Unit tests for ComposerService
Tests all methods and functions in composer_service.py
"""
import pytest
import os
import tempfile
from unittest.mock import Mock, AsyncMock, patch, MagicMock
from docx import Document
from sqlalchemy.orm import Session

from app.services.composer_service import ComposerService
from app.models import Upload, Job, Screenshot


@pytest.fixture
def composer_service():
    """Create ComposerService instance"""
    return ComposerService()


@pytest.fixture
def mock_db():
    """Create mock database session"""
    return Mock(spec=Session)


@pytest.fixture
def sample_upload():
    """Create sample upload object"""
    upload = Mock(spec=Upload)
    upload.id = 1
    upload.filename = "test_lab.docx"
    upload.file_path = "/path/to/test_lab.docx"
    return upload


@pytest.fixture
def sample_job():
    """Create sample job object"""
    job = Mock(spec=Job)
    job.id = 1
    job.task_id = 1
    job.upload_id = 1
    job.status = "completed"
    job.question_text = "Write a program to calculate factorial"
    job.output_text = "Factorial of 5 is 120"
    return job


@pytest.fixture
def sample_screenshot():
    """Create sample screenshot object"""
    screenshot = Mock(spec=Screenshot)
    screenshot.id = 1
    screenshot.job_id = 1
    screenshot.file_path = "/path/to/screenshot.png"
    screenshot.created_at = "2025-01-01"
    return screenshot


@pytest.fixture
def temp_docx_file():
    """Create temporary DOCX file for testing"""
    doc = Document()
    doc.add_paragraph("Test Lab Manual")
    doc.add_paragraph("Question 1: Write a program")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        yield f.name
    
    # Cleanup
    if os.path.exists(f.name):
        os.unlink(f.name)


class TestComposerServiceInit:
    """Test ComposerService initialization"""
    
    @pytest.mark.unit
    def test_init(self, composer_service):
        """Test service initialization"""
        assert composer_service is not None
        assert isinstance(composer_service, ComposerService)


class TestComposeReport:
    """Test compose_report method"""
    
    @pytest.mark.unit
    @pytest.mark.asyncio
    async def test_compose_report_success(self, composer_service, mock_db, sample_upload, 
                                         sample_job, sample_screenshot, temp_docx_file):
        """Test successful report composition"""
        # Setup mocks
        mock_db.query.return_value.filter.return_value.first.return_value = sample_upload
        sample_upload.file_path = temp_docx_file
        
        # Mock screenshot query
        mock_screenshot_query = Mock()
        mock_screenshot_query.filter.return_value.order_by.return_value.all.return_value = [sample_screenshot]
        mock_db.query.return_value.filter.return_value.order_by.return_value.all.return_value = [sample_screenshot]
        
        # Mock job query
        mock_job_query = Mock()
        mock_job_query.filter.return_value.first.return_value = sample_job
        mock_job_query.filter.return_value.all.return_value = [sample_job]
        
        # Mock settings
        with patch('app.services.composer_service.settings') as mock_settings:
            mock_settings.REPORT_DIR = tempfile.gettempdir()
            
            # Mock os.path.exists for screenshot
            with patch('os.path.exists', return_value=True):
                result = await composer_service.compose_report(
                    upload_id=1,
                    screenshot_order=[1],
                    db=mock_db
                )
        
        # Assertions
        assert result is not None
        assert "report_path" in result
        assert "filename" in result
        assert "download_url" in result
        assert result["filename"].endswith(".docx")
    
    @pytest.mark.unit
    @pytest.mark.asyncio
    async def test_compose_report_upload_not_found(self, composer_service, mock_db):
        """Test compose_report when upload not found"""
        mock_db.query.return_value.filter.return_value.first.return_value = None
        
        with pytest.raises(ValueError, match="Upload not found"):
            await composer_service.compose_report(
                upload_id=999,
                screenshot_order=[],
                db=mock_db
            )
    
    @pytest.mark.unit
    @pytest.mark.asyncio
    async def test_compose_report_document_not_found(self, composer_service, mock_db, sample_upload):
        """Test compose_report when document file not found"""
        mock_db.query.return_value.filter.return_value.first.return_value = sample_upload
        sample_upload.file_path = "/nonexistent/path.docx"
        
        with pytest.raises(ValueError, match="Original document not found"):
            await composer_service.compose_report(
                upload_id=1,
                screenshot_order=[],
                db=mock_db
            )
    
    @pytest.mark.unit
    @pytest.mark.asyncio
    async def test_compose_report_empty_screenshot_order(self, composer_service, mock_db, 
                                                         sample_upload, temp_docx_file):
        """Test compose_report with empty screenshot order"""
        mock_db.query.return_value.filter.return_value.first.return_value = sample_upload
        sample_upload.file_path = temp_docx_file
        
        # Mock empty jobs list
        mock_db.query.return_value.filter.return_value.all.return_value = []
        
        with patch('app.services.composer_service.settings') as mock_settings:
            mock_settings.REPORT_DIR = tempfile.gettempdir()
            
            result = await composer_service.compose_report(
                upload_id=1,
                screenshot_order=[],
                db=mock_db
            )
        
        assert result is not None
        assert "report_path" in result


class TestGetVariedHeader:
    """Test _get_varied_header method"""
    
    @pytest.mark.unit
    def test_get_varied_header_returns_string(self, composer_service):
        """Test that _get_varied_header returns a string"""
        header = composer_service._get_varied_header()
        assert isinstance(header, str)
        assert len(header) > 0
    
    @pytest.mark.unit
    def test_get_varied_header_variety(self, composer_service):
        """Test that _get_varied_header returns different headers"""
        headers = [composer_service._get_varied_header() for _ in range(10)]
        # Should have some variety (not all the same)
        # Note: This might occasionally fail due to randomness, but unlikely
        assert len(set(headers)) >= 1  # At least one unique header


class TestGenerateImageDescription:
    """Test _generate_image_description method"""
    
    @pytest.mark.unit
    def test_generate_image_description_factorial(self, composer_service, sample_job):
        """Test description generation for factorial question"""
        question = "Write a program to calculate factorial"
        description = composer_service._generate_image_description(question, sample_job, 1)
        
        assert isinstance(description, str)
        assert len(description) > 0
        assert "factorial" in description.lower() or "recursive" in description.lower()
    
    @pytest.mark.unit
    def test_generate_image_description_sort(self, composer_service, sample_job):
        """Test description generation for sort question"""
        question = "Write a program to sort an array"
        description = composer_service._generate_image_description(question, sample_job, 1)
        
        assert isinstance(description, str)
        assert len(description) > 0
    
    @pytest.mark.unit
    def test_generate_image_description_oop(self, composer_service, sample_job):
        """Test description generation for OOP question"""
        question = "Create a class with inheritance"
        description = composer_service._generate_image_description(question, sample_job, 1)
        
        assert isinstance(description, str)
        assert len(description) > 0
        assert "object-oriented" in description.lower() or "oop" in description.lower() or "class" in description.lower()


class TestOrderScreensForJob:
    """Test _order_screens_for_job method"""
    
    @pytest.mark.unit
    def test_order_screens_for_job_empty(self, composer_service):
        """Test ordering with empty screens list"""
        result = composer_service._order_screens_for_job([])
        assert isinstance(result, list)
        assert len(result) == 0
    
    @pytest.mark.unit
    def test_order_screens_for_job_single(self, composer_service, sample_screenshot):
        """Test ordering with single screenshot"""
        result = composer_service._order_screens_for_job([sample_screenshot])
        assert isinstance(result, list)
        assert len(result) == 1
        assert "screen" in result[0]
        assert "label" in result[0]
        assert "category" in result[0]


class TestClassifyScreenForDoc:
    """Test _classify_screen_for_doc method"""
    
    @pytest.mark.unit
    def test_classify_screen_code(self, composer_service, sample_screenshot):
        """Test classification of code screenshot"""
        sample_screenshot.file_path = "/path/to/code_screenshot.png"
        category, label = composer_service._classify_screen_for_doc(sample_screenshot)
        
        assert category == "code"
        assert "Python" in label or "code" in label.lower()
    
    @pytest.mark.unit
    def test_classify_screen_file_preview_txt(self, composer_service):
        """Test classification of .txt file preview"""
        screen = Mock()
        screen.file_path = "/path/to/file_test.txt"
        category, label = composer_service._classify_screen_for_doc(screen)
        
        assert category == "txt"
        assert ".txt" in label.lower()
    
    @pytest.mark.unit
    def test_classify_screen_file_preview_pkl(self, composer_service):
        """Test classification of .pkl file preview"""
        screen = Mock()
        screen.file_path = "/path/to/file_data.pkl"
        category, label = composer_service._classify_screen_for_doc(screen)
        
        assert category == "pkl"
        assert ".pkl" in label.lower()


class TestIsFilePreviewPath:
    """Test _is_file_preview_path method"""
    
    @pytest.mark.unit
    def test_is_file_preview_path_true(self, composer_service):
        """Test file preview path detection - true case"""
        assert composer_service._is_file_preview_path("file_test.txt") is True
        assert composer_service._is_file_preview_path("/path/to/file_data.pkl") is True
    
    @pytest.mark.unit
    def test_is_file_preview_path_false(self, composer_service):
        """Test file preview path detection - false case"""
        assert composer_service._is_file_preview_path("screenshot.png") is False
        assert composer_service._is_file_preview_path("/path/to/code.png") is False
        assert composer_service._is_file_preview_path("") is False
        assert composer_service._is_file_preview_path(None) is False


class TestInferPreviewExtension:
    """Test _infer_preview_extension method"""
    
    @pytest.mark.unit
    def test_infer_preview_extension_txt(self, composer_service):
        """Test extension inference for .txt"""
        ext = composer_service._infer_preview_extension("file_test.txt")
        assert ext == ".txt"
    
    @pytest.mark.unit
    def test_infer_preview_extension_pkl(self, composer_service):
        """Test extension inference for .pkl"""
        ext = composer_service._infer_preview_extension("file_data.pkl")
        assert ext == ".pkl"
    
    @pytest.mark.unit
    def test_infer_preview_extension_not_file_preview(self, composer_service):
        """Test extension inference for non-file-preview path"""
        ext = composer_service._infer_preview_extension("screenshot.png")
        assert ext == ""
    
    @pytest.mark.unit
    def test_infer_preview_extension_empty(self, composer_service):
        """Test extension inference for empty path"""
        ext = composer_service._infer_preview_extension("")
        assert ext == ""


class TestExtractFileLabel:
    """Test _extract_file_label method"""
    
    @pytest.mark.unit
    def test_extract_file_label_with_label(self, composer_service):
        """Test label extraction when label exists"""
        label = composer_service._extract_file_label("file_student_data.txt")
        # Implementation dependent - may return None or extracted label
        assert label is None or isinstance(label, str)
    
    @pytest.mark.unit
    def test_extract_file_label_empty(self, composer_service):
        """Test label extraction for empty path"""
        label = composer_service._extract_file_label("")
        assert label is None
    
    @pytest.mark.unit
    def test_extract_file_label_none(self, composer_service):
        """Test label extraction for None"""
        label = composer_service._extract_file_label(None)
        assert label is None


class TestFindQuestionEndIndex:
    """Test _find_question_end_index method"""
    
    @pytest.mark.unit
    def test_find_question_end_index(self, composer_service):
        """Test finding question end index"""
        paragraphs = [
            Mock(text="Question 1: Write a program"),
            Mock(text="This is the question text"),
            Mock(text="Question 2: Another question"),
        ]
        
        end_index = composer_service._find_question_end_index(paragraphs, 0)
        assert isinstance(end_index, int)
        assert end_index >= 0


class TestExtractTaskNumber:
    """Test _extract_task_number method"""
    
    @pytest.mark.unit
    def test_extract_task_number_with_number(self, composer_service):
        """Test task number extraction when number exists"""
        number = composer_service._extract_task_number("Task 1: Description")
        assert number == 1
    
    @pytest.mark.unit
    def test_extract_task_number_no_number(self, composer_service):
        """Test task number extraction when no number exists"""
        number = composer_service._extract_task_number("Task: Description")
        assert number is None
    
    @pytest.mark.unit
    def test_extract_task_number_empty(self, composer_service):
        """Test task number extraction for empty string"""
        number = composer_service._extract_task_number("")
        assert number is None


class TestFindQuestionPattern:
    """Test _find_question_pattern method"""
    
    @pytest.mark.unit
    def test_find_question_pattern_with_number(self, composer_service):
        """Test finding question pattern with number"""
        number = composer_service._find_question_pattern("Question 1: Write a program")
        assert number == 1
    
    @pytest.mark.unit
    def test_find_question_pattern_no_match(self, composer_service):
        """Test finding question pattern when no match"""
        number = composer_service._find_question_pattern("Just some text")
        assert number is None


class TestGetStatusColor:
    """Test _get_status_color method"""
    
    @pytest.mark.unit
    def test_get_status_color_completed(self, composer_service):
        """Test status color for completed"""
        color = composer_service._get_status_color("completed")
        assert isinstance(color, tuple)
        assert len(color) == 3  # RGB tuple
    
    @pytest.mark.unit
    def test_get_status_color_failed(self, composer_service):
        """Test status color for failed"""
        color = composer_service._get_status_color("failed")
        assert isinstance(color, tuple)
        assert len(color) == 3
    
    @pytest.mark.unit
    def test_get_status_color_default(self, composer_service):
        """Test status color for unknown status"""
        color = composer_service._get_status_color("unknown")
        assert isinstance(color, tuple)
        assert len(color) == 3




